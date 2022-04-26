import docx2txt
import re
import docx


class WordHandler():
    def __init__(self, path):
        self.text = docx2txt.process(path)

    def get_tag(self):
        pattern = r"\[(.*?)\]"
        r = re.findall(pattern, self.text)
        return r

    def get_summary(self, tags):
        # TODO : Implement summary funciton
        total = len(tags)
        pics = tags.count("PIC")

        print(total, pics)

def test():
    path = "./demo.docx"
    handler = WordHandler(path)
    tags = (handler.get_tag())

    print(tags)

    handler.get_summary(tags)

class NewHandler():
    def __init__(self, path):
        self.doc = docx.Document(path)
    
    def get_tag(self):
        tags = []
        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font.highlight_color:
                    tags.append(run.text)
                elif run.bold:
                    tags.append(f"PNT: {run.text}")    
        return tags

def new():
    import docx
    doc = docx.Document("./demo2.docx")
    tags = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.highlight_color:
                tags.append(run.text)
            elif run.bold:
                print(run.text)
                tags.append(f"PNT: {run.text}")
    print(tags)
if __name__ == "__main__":
    # test()
    new()